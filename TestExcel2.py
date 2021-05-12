# -*- coding: utf-8 -*-
"""
Created on Wed May 12 10:15:47 2021

@author: User
"""
import pandas as pd
import os
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
#excel表格路径
data = pd.read_excel("E:/Desktop/test.xlsx") 
#word模板路径
docxPath = "E:/Desktop/table.docx"
savePath = "E:/Desktop/TestExcel/"
picPath =  "E:/Desktop/pic/"

'''
合并文件函数 
将sub的word文档内容合并到master中

'''
def combine_docx(master, sub):
    if not os.path.exists(sub):#待合并文件必须存在
        print(sub)
       
        return 1

    if not master.endswith('.docx'):#主文件必须是docx格式（可以不存在）
        return 2

    sub_docx = sub
   

    if os.path.exists(master):
        doc_master = Document(master)
        #doc_master.add_page_break()
        cp = Composer(doc_master)
        cp.append(Document(sub_docx))
    else:
        #master不存在，则sub直接给master
        doc_master = Document(sub_docx)

    doc_master.save(master)
    return True


count = 0

for i in range(data.shape[0]):
    tpl = DocxTemplate(docxPath)
    if(i%2==0):
        textline=[]
    for j in range(data.shape[1]):
        textline.append(data.iloc[i][j])
   # print(textline)
    
    if(i%2!=0 and i!=0):
        if(i<10):
            textline[3]='0'+str(textline[3])
            textline[9]='0'+str(textline[9])
          
        
        context = {
            "name1": textline[1],
            "examno1" : textline[4],
            "idno1" : textline[2],
            "no1" : textline[3],
            "addr1" : textline[5], 
            "name2": textline[7],
            "examno2" : textline[10],
            "idno2" : textline[8],
            "no2" : textline[9],
            "addr2" : textline[11],  }
        
        
        tpl.render(context)
        #在word模板中插入一个类似于模板的照片，1.jpeg 与2.jpeg分别是插入的模板图片的名字
        
        tpl.replace_pic('1.jpeg', picPath + textline[1] + textline[2] +".jpg")
        tpl.replace_pic('2.jpeg', picPath + textline[7] + textline[8] +".jpg")
        
        tpl.save(savePath + "{}.docx".format(count + 1))
        if(count!=0):
            flag = combine_docx(savePath+"1.docx", savePath +str(count + 1)+".docx")
            print(flag)
        count = count +1
#需求是 将文件名字改为 教室的地址
os.renames(savePath+"1.docx", savePath + textline[11]+".docx")

